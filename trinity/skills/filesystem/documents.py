"""
Trinity Skill — Document Handler.
Handles reading and creating .docx, .xlsx, .pdf, .pptx files.
"""

import structlog
from pathlib import Path
from trinity.skills.base import BaseSkill, SkillResult

logger = structlog.get_logger(__name__)


class DocumentHandler(BaseSkill):
    """Handle document creation and editing for Office formats."""

    async def execute(self, entities: dict, context: dict | None = None) -> SkillResult:
        """Execute document operation."""
        path_str = entities.get("path", "")
        path = Path(self._resolve_path(path_str))
        suffix = path.suffix.lower()

        if suffix == ".docx":
            return await self._handle_docx(entities, path)
        elif suffix == ".xlsx":
            return await self._handle_xlsx(entities, path)
        elif suffix == ".pdf":
            return await self._handle_pdf(entities, path)
        else:
            return self._error(f"Unsupported document format: {suffix}")

    async def _handle_docx(self, entities: dict, path: Path) -> SkillResult:
        """Handle Word document operations."""
        try:
            from docx import Document

            if not path.exists():
                # Create new document
                doc = Document()
                content = entities.get("content", "")
                if content:
                    for paragraph in content.split("\n"):
                        doc.add_paragraph(paragraph)
                doc.save(str(path))
                return self._success(f"Created Word document: {path.name}")

            # Read existing document
            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs]
            text = "\n".join(paragraphs)
            return self._success(f"Content of {path.name}:\n\n{text}",
                                data={"paragraphs": len(paragraphs)})

        except ImportError:
            return self._error("python-docx not installed. Run: pip install python-docx")

    async def _handle_xlsx(self, entities: dict, path: Path) -> SkillResult:
        """Handle Excel spreadsheet operations."""
        try:
            from openpyxl import Workbook, load_workbook

            if not path.exists():
                wb = Workbook()
                ws = wb.active
                ws.title = "Sheet1"
                wb.save(str(path))
                return self._success(f"Created spreadsheet: {path.name}")

            wb = load_workbook(str(path), read_only=True)
            result = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                result.append(f"--- {sheet_name} ---")
                for row in ws.iter_rows(values_only=True):
                    result.append(" | ".join(str(c) if c else "" for c in row))
            wb.close()

            return self._success("\n".join(result))

        except ImportError:
            return self._error("openpyxl not installed. Run: pip install openpyxl")

    async def _handle_pdf(self, entities: dict, path: Path) -> SkillResult:
        """Handle PDF operations (read only, create new)."""
        try:
            import fitz  # PyMuPDF

            if not path.exists():
                # Create a simple PDF
                doc = fitz.open()
                page = doc.new_page()
                content = entities.get("content", "Created by Trinity")
                page.insert_text((72, 72), content)
                doc.save(str(path))
                doc.close()
                return self._success(f"Created PDF: {path.name}")

            # Read existing PDF
            doc = fitz.open(str(path))
            text = ""
            for i, page in enumerate(doc):
                text += f"\n--- Page {i+1} ---\n"
                text += page.get_text()
            doc.close()

            return self._success(f"Content of {path.name} ({doc.page_count} pages):\n{text[:10000]}")

        except ImportError:
            return self._error("PyMuPDF not installed. Run: pip install PyMuPDF")
